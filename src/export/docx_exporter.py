"""Word document exporter using python-docx."""

from pathlib import Path
from typing import Optional

from src.export.base import BaseExporter
from src.models.paper import PaperDraft
from src.models.session import ResearchSession


class DocxExporter(BaseExporter):
    """Export papers to Word (.docx) format."""

    format_name = "docx"
    file_extension = "docx"

    async def export(
        self,
        draft: PaperDraft,
        session: ResearchSession,
        output_dir: str,
    ) -> Optional[str]:
        """Export the paper to Word format."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE

            output_path = self._get_output_path(session, output_dir)
            self._logger.info(f"Exporting to DOCX: {output_path}")

            # Create document
            doc = Document()

            # Set up styles
            self._setup_styles(doc)

            # Title page
            self._add_title_page(doc, session, draft)

            # Abstract
            if "abstract" in draft.sections:
                doc.add_page_break()
                self._add_section(doc, "Abstract", draft.sections["abstract"])

            # Main sections
            doc.add_page_break()

            section_order = [
                ("introduction", "Introduction"),
                ("literature_review", "Literature Review"),
                ("theoretical_framework", "Theoretical Framework"),
                ("methodology", "Methodology"),
                ("analysis", "Analysis and Findings"),
                ("discussion", "Discussion"),
                ("conclusion", "Conclusion"),
                ("references", "References"),
            ]

            for key, title in section_order:
                content = draft.sections.get(key, "")
                if content:
                    self._add_section(doc, title, content)

            # Save document
            doc.save(str(output_path))

            self._logger.info(f"DOCX exported successfully: {output_path}")
            return str(output_path)

        except ImportError:
            self._logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            self._logger.error(f"DOCX export failed: {e}")
            return None

    def _setup_styles(self, doc) -> None:
        """Set up document styles."""
        from docx.shared import Pt
        from docx.enum.style import WD_STYLE_TYPE

        # Modify Normal style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 2.0

        # Heading 1 style
        if 'Heading 1' in doc.styles:
            h1 = doc.styles['Heading 1']
            h1.font.name = 'Times New Roman'
            h1.font.size = Pt(14)
            h1.font.bold = True

        # Heading 2 style
        if 'Heading 2' in doc.styles:
            h2 = doc.styles['Heading 2']
            h2.font.name = 'Times New Roman'
            h2.font.size = Pt(12)
            h2.font.bold = True

    def _add_title_page(self, doc, session: ResearchSession, draft: PaperDraft) -> None:
        """Add title page to the document."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Add some space at the top
        for _ in range(8):
            doc.add_paragraph()

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(session.title)
        title_run.bold = True
        title_run.font.size = Pt(18)

        # Spacing
        for _ in range(4):
            doc.add_paragraph()

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Citation Style: {session.citation_style.value}")

        meta2 = doc.add_paragraph()
        meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta2.add_run(f"Word Count: {draft.word_count}")

    def _add_section(self, doc, title: str, content: str) -> None:
        """Add a section to the document."""
        from docx.shared import Inches

        # Section heading
        doc.add_heading(title, level=1)

        # Section content - split into paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = doc.add_paragraph()
                # Handle basic formatting
                self._add_formatted_text(para, para_text)
                para.paragraph_format.first_line_indent = Inches(0.5)

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with basic formatting to a paragraph."""
        import re

        # Simple bold and italic parsing
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)
